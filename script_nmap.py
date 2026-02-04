
import paramiko
import argparse
import logging
import time
import math
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm
from docx import Document

def setup_logging(network):
    log_filename = f"scan_log_{network.replace('/', '_')}.txt"
    logging.basicConfig(
        filename=log_filename,
        level=logging.DEBUG,
        format="%(asctime)s - %(levelname)s - %(message)s"
    )

def execute_remote_command_as_root(client, command, sudo_password, debug=False):
    if debug:
        logging.info("Eseguendo comando con sudo: %s", command)
    command = f"echo {sudo_password} | sudo -S {command}"
    stdin, stdout, stderr = client.exec_command(command)
    error_output = stderr.read().decode()
    if error_output:
        logging.error("Errore comando '%s':\n%s", command, error_output)
    return stdout.read().decode() + error_output

def execute_udp_nmap(client, host, udp_ports, sudo_password, debug=False):
    udp_ports_str = ",".join(udp_ports)
    command = f"nmap -sU -sV -p{udp_ports_str} -T4 -Pn {host}"
    if debug:
        logging.info("Nmap UDP -sV %s porte: %s", host, udp_ports_str)
    return execute_remote_command_as_root(client, command, sudo_password, debug)

def main():
    parser = argparse.ArgumentParser(description="Scansione TCP/UDP dinamica con output DOCX.")
    parser.add_argument("remote_host", type=str)
    parser.add_argument("username", type=str)
    parser.add_argument("password", type=str)
    parser.add_argument("network", type=str)
    parser.add_argument("--max-hours", type=int, default=12)
    parser.add_argument("--exclude", type=str)
    parser.add_argument("-d", "--debug", action="store_true")
    args = parser.parse_args()

    remote_host, username, password, network = args.remote_host, args.username, args.password, args.network
    max_hours, debug = args.max_hours, args.debug

    setup_logging(network)

    client = paramiko.SSHClient()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    client.connect(remote_host, username=username, password=password)

    try:
        ping_command = f"nmap -n -sn -T4 -PS3389,22,80,443 -PA3389,22,80,443 -PU137,40125 {network}"
        result = execute_remote_command_as_root(client, ping_command, password, debug)
        logging.info("Ping scan:\n%s", result.strip())

        active_hosts = []
        exclude_ips = [ip.strip() for ip in args.exclude.split(",")] if args.exclude else []
        for line in result.splitlines():
            if "Nmap scan report for" in line:
                ip = line.split()[-1]
                if ip not in exclude_ips:
                    active_hosts.append(ip)

        total_time_budget = max_hours * 3600
        seconds_per_host = total_time_budget / max(len(active_hosts), 1)
        max_threads = max(1, min(10, math.ceil(len(active_hosts) * 1200 / total_time_budget)))

        results_filename = f"nmap_results_{network.replace('/', '_')}.txt"
        docx_filename = results_filename.replace(".txt", ".docx")
        start_time = time.time()

        doc = Document()
        doc.add_heading(f"Report di Scansione Rete: {network}", level=1)

        def scan_host(host):
            if seconds_per_host > 1800:
                tcp_cmd = f"nmap -sS -p- -T4 -Pn {host}"
                udp_cmd = f"masscan -pU:0-65535 {host} --rate 1000"
            elif seconds_per_host > 900:
                tcp_cmd = f"nmap -sS -p- -T4 -Pn {host}"
                udp_cmd = f"masscan -pU:0-65535 {host} --rate 500"
            elif seconds_per_host > 300:
                tcp_cmd = f"nmap -sS -p- -T3 -Pn {host}"
                udp_cmd = f"masscan -pU:0-65535 {host} --rate 250"
            else:
                tcp_cmd = f"nmap -sS -p- -T3 -Pn {host}"
                udp_cmd = f"masscan -pU:0-65535 {host} --rate 100"

            tcp_result = execute_remote_command_as_root(client, tcp_cmd, password, debug)
            tcp_ports = [line.split("/")[0].strip() for line in tcp_result.splitlines() if "/tcp" in line and "open" in line]

            if tcp_ports:
                ports_str = ",".join(tcp_ports)
                tcp_sV_cmd = f"nmap -sS -sV -Pn -p{ports_str} {host}"
                tcp_sV_result = execute_remote_command_as_root(client, tcp_sV_cmd, password, debug)
                doc.add_heading(f"Risultati scansione TCP per {host}", level=2)
                doc.add_paragraph(tcp_sV_result.strip())
                with open(results_filename, "a") as f:
                    f.write(f"===== TCP -sV {host} =====\n{tcp_sV_result}\n")

            udp_result = execute_remote_command_as_root(client, udp_cmd, password, debug)
            udp_ports = []
            for line in udp_result.splitlines():
                if "U:" in line:
                    for part in line.split():
                        if part.startswith("U:"):
                            udp_ports.append(part.split(":")[1])

            if udp_ports:
                udp_nmap_result = execute_udp_nmap(client, host, udp_ports, password, debug)
                doc.add_heading(f"Risultati scansione UDP per {host}", level=2)
                doc.add_paragraph(udp_nmap_result.strip())
                with open(results_filename, "a") as f:
                    f.write(f"===== UDP -sV {host} =====\n{udp_nmap_result}\n")

        with ThreadPoolExecutor(max_workers=max_threads) as executor:
            futures = [executor.submit(scan_host, host) for host in active_hosts]
            for future in tqdm(as_completed(futures), total=len(futures), desc='Scansione host'):
                future.result()

        elapsed = time.time() - start_time
        doc.save(docx_filename)
        print(f"✅ Scansione completata in {elapsed / 60:.2f} minuti")

    finally:
        client.close()

if __name__ == "__main__":
    main()
